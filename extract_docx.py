import docx
import sys

def extract_docx(file_path):
    doc = docx.Document(file_path)
    with open("extract.txt", "w", encoding="utf-8") as f:
        for p in doc.paragraphs:
            if p.text.strip():
                f.write(p.text + "\n")
        
        f.write("\n--- TABLES ---\n")
        for i, table in enumerate(doc.tables):
            f.write(f"Table {i+1}:\n")
            for row in table.rows:
                f.write(" | ".join([c.text.replace("\n", " ").strip() for c in row.cells]) + "\n")
            f.write("\n")

if __name__ == "__main__":
    extract_docx("docs/SCNV_Gap_Analysis_Enhancement_Plan.docx")
